# Imports
import os
import sys
import subprocess
cwd = os.getcwd()
p = cwd + "\\Storage"
# Check whether the specified path exists or not
dependencies = ['tk', 'tkhtmlview', 'nltk', 'sentence_transformers', 'matplotlib', 'scipy', 'scikit-learn', 'python-docx', 'datetime', 'pywin32']
if not os.path.exists(p):
    setup = input("First time use, install dependancies? (y/n) :")
    if setup == 'y':
        for k in dependencies:
            subprocess.check_call([sys.executable, '-m', 'pip', 'install', k])
        import nltk
        nltk.download('punkt')
        nltk.download("stopwords")
        nltk.download('wordnet')

from tkinter import *
import tkinter as tk
from tkinter.filedialog import askdirectory
import pickle
import sqlite3
from tkhtmlview import HTMLLabel
import nltk
from sentence_transformers import SentenceTransformer, util
import matplotlib.pyplot as plt
from scipy.cluster.hierarchy import dendrogram, linkage
from sklearn.cluster import AgglomerativeClustering
import docx
from datetime import datetime
import shutil
import win32com.client
import ctypes


# Functions


class Biblio:

    def __init__(self, window):
        cwd = os.getcwd()
        self.p = cwd + "\\Storage"
        self.blocs = self.import_dict(self.p, 'blocs')
        self.tag_list = self.build_tag_list(self.blocs)
        self.plan = self.import_dict(self.p, 'plan')
        self.zotero = self.import_dict(self.p, 'Zotero_data')
        self.tagging = 0
        self.noting = 0
        self.accepted_order = [k for k in range(6)]
        self.window = window

        self.var = IntVar()

        self.merge_var = IntVar()

        self.save_var = IntVar()
        
        self.mothermenu = tk.Frame(window, bg=BG_COLOR)
        self.mothermenu.grid(row=0, column=0, sticky=tk.W + tk.E)

        tk.Button(
            self.mothermenu,
            text='Locate Zotero',
            height=1,
            width=10,
            command=self.locate_zotero).grid(row=0, column=0, sticky=tk.W)
        
        self.zotero_path = tk.StringVar()
        self.zotero_path.set(
            self.zotero['path'] if self.zotero else 'Not yet connected to Zotero.')
        self.show_zotero_path = tk.Label(
            self.mothermenu, textvariable=self.zotero_path)
        self.show_zotero_path.grid(row=0, column=1, sticky=tk.W+tk.E)
        
        display1 = tk.Frame(window, bg=BG_COLOR)
        display1.grid(row=1, column=0)

        display1arrow = tk.Frame(display1, bg=BG_COLOR, bd=10)
        display1arrow.grid(row=0, column=0)

        tk.Button(
            display1arrow,
            text='^',
            height=1,
            width=3,
            command=self.move_up_plan).grid(row=0, column=1)

        tk.Button(
            display1arrow,
            text='v',
            height=1,
            width=3,
            command=self.move_down_plan).grid(row=2, column=1)

        tk.Button(
            display1arrow,
            text='<',
            height=1,
            width=3,
            command=self.move_left_plan).grid(row=1, column=0)

        tk.Button(
            display1arrow,
            text='>',
            height=1,
            width=3,
            command=self.move_right_plan).grid(row=1, column=2)

        display1plan = tk.Frame(display1, bg=BG_COLOR, bd=10)
        display1plan.grid(row=0, column=1, sticky=tk.W + tk.E)

        self.plan_listbox = Listbox(
            display1plan,
            height=30,
            width=50,
            selectmode=EXTENDED)
        self.plan_listbox.grid(row=0, column=0)
        new = self.build_plan()
        for k in range(len(new)):
            self.plan_listbox.insert(k, new[k])
        self.plan_listbox.bind('<<ListboxSelect>>', self.blocs_filter_plan)

        display1planmenu = tk.Frame(display1, bg=BG_COLOR, bd=10)
        display1planmenu.grid(row=1, column=1, sticky=tk.W + tk.E)

        display1planmenu1 = tk.Frame(display1planmenu, bg=BG_COLOR)
        display1planmenu1.grid(row=0, column=0)
        tk.Button(
            display1planmenu1,
            text='Add here',
            height=1,
            width=7,
            command=self.add_plan).grid(row=0, column=0)

        display1planmenu11 = tk.Frame(
            display1planmenu1)
        display1planmenu11.grid()
        tk.Button(
            display1planmenu11,
            text='-> add lower',
            height=1,
            width=11,
            command=self.add_plan_low).grid(row=1, column=0)

        tk.Button(
            display1planmenu11,
            text='add upper <-',
            height=1,
            width=11,
            command=self.add_plan_high).grid(row=1, column=1)
            
        display1planmenu2 = tk.Frame(display1planmenu)
        display1planmenu2.grid(row=0, column=1)
        
        tk.Button(
            display1planmenu2,
            text='Delete',
            height=1,
            width=7,
            command=self.delete_plan).grid(row=0, column=0, sticky=tk.E)

        tk.Button(
            display1planmenu2,
            text='Edit',
            height=1,
            width=7,
            command=self.edit_plan).grid(row=1, column=0, sticky=tk.E)

        display1source = tk.Frame(display1, bg=BG_COLOR, bd=10)
        display1source.grid(row=0, column=2)

        self.source_listbox = Listbox(
            display1source,
            height=30,
            width=20,
            selectmode=EXTENDED)
        self.source_listbox.grid(row=0, column=0)
        for k in unique(self.blocs["source"]):
            self.source_listbox.insert(END, k[0])
        self.source_listbox.bind('<<ListboxSelect>>', self.blocs_filter_sources)

        display1sourcemenu = tk.Frame(display1, bg=BG_COLOR, bd=10)
        display1sourcemenu.grid(row=1, column=2)

        importzoterohighlight = tk.Frame(
            display1sourcemenu, bg='red', bd=1)
        importzoterohighlight.grid(row=0, column=0, sticky=tk.W + tk.E)
        tk.Button(
            importzoterohighlight,
            text='Import Zot',
            height=1,
            width=9,
            command=self.add_to_blocs).grid(row=0, column=0, sticky=tk.W + tk.E)
            
        tk.Button(
            display1sourcemenu,
            text='Info',
            height=1,
            width=9,
            command=self.send_key).grid(row=0, column=1, sticky=tk.W + tk.E)

        tk.Button(
            display1sourcemenu,
            text='Remove',
            height=1,
            width=7,
            command=self.delete_article).grid(row=1, column=0, sticky=tk.W + tk.E)

        display1blocs = tk.Frame(display1, bg=BG_COLOR, bd=10)
        display1blocs.grid(row=0, column=3)

        self.blocs_listbox = Listbox(
            display1blocs,
            height=30,
            width=60,
            selectmode=EXTENDED)
        self.blocs_listbox.grid(row=0, column=0)
        self.blocs_listbox.bind('<<ListboxSelect>>', self.read_blocs)

        display1blocsoptions = tk.Frame(display1, bg=BG_COLOR, bd=10)
        display1blocsoptions.grid(row=1, column=3)

        tk.Button(
            display1blocsoptions,
            text='Tagging',
            height=1,
            width=10,
            command=self.tag_blocs).grid(row=0, column=0)

        tk.Checkbutton(
            display1blocsoptions,
            text='Merge ?',
            bg=BG_COLOR,
            variable=self.merge_var).grid(row=0, column=1)

        display1shell = tk.Frame(display1, bg=BG_COLOR, bd=10)
        display1shell.grid(row=0, column=4, sticky=tk.N + tk.S)

        self.shell_label = Label(display1shell, text="Shell :", bg=BG_COLOR)
        self.shell_label.grid(row=0, column=0, sticky=tk.W)

        self.shell_text = Text(
            display1shell,
            #height=23,
            width=60,
            exportselection=False,
            wrap=WORD,
            font=('Calibri', 12))
        self.shell_text.grid(row=1, column=0, sticky=tk.S)

        display1shelloptions = tk.Frame(display1, bg=BG_COLOR, bd=10)
        display1shelloptions.grid(row=1, column=4, sticky=tk.W)

        self.search_text = Text(
            display1shelloptions,
            height=1,
            width=40)
        self.search_text.grid(row=0, column=0)

        self.search_but = Button(
            display1shelloptions,
            text='Search',
            height=1,
            width=7,
            command=self.blocs_filter_search)
        self.search_but.grid(row=0, column=1, sticky=tk.E)

        self.topics_but = Button(
            display1shelloptions,
            text='Topics',
            height=1,
            width=7,
            command=self.blocs_main_subjects)
        self.topics_but.grid(row=0, column=2, sticky=tk.E)
        
        display2 = tk.Frame(window, bg=BG_COLOR)
        display2.grid(row=2, column=0, sticky=tk.W + tk.E)

        display2note = tk.Frame(display2, bg=BG_COLOR, bd=10)
        display2note.grid(row=0, column=0, sticky=tk.W)

        self.notes_text = Text(
            display2note,
            height=11,
            width=150,
            wrap=WORD,
            font=('Calibri',12))
        self.notes_text.grid(row=0, column=0)

        display2menu = tk.Frame(display2, bg=BG_COLOR, bd=10)
        display2menu.grid(row=0, column=1, sticky=tk.W + tk.E)

        self.take_note_but = tk.Button(
            display2menu,
            text='Take notes',
            height=1,
            width=10,
            command=self.edit_notes_from_plan)
        self.take_note_but.grid(row=0, column=0)

        self.ref_but = Button(
            display2menu,
            text='Link Ref',
            height=1,
            width=10,
            command=self.insert_ref)
        self.ref_but.grid(row=1, column=0)

        self.save_note_but = Button(
            display2menu,
            text='Save',
            height=1,
            width=10,
            command=lambda: self.save_var.set(1))
        self.save_note_but.grid(row=2, column=0)

        self.export_but = Button(
            display2menu,
            text='Export all',
            height=1,
            width=10,
            command=self.export_all)
        self.export_but.grid(row=3, column=1)

        self.backup_but = Button(
            display2menu,
            text='Backup',
            height=1,
            width=10,
            command=self.backup)
        self.backup_but.grid(row=4, column=1)

    # Dictionary management

    def import_dict(self, path, name):
        with open(path + '\\' + name + '.pkl', 'rb') as f:
            return pickle.load(f)

    def save_dict(self, path, name, d):
        with open(path + '\\' + name + '.pkl', 'wb') as f:
            pickle.dump(d, f)

    # Tag plan management

    def build_tag_list(self, blocs):
        tl = []
        for i in blocs["tag"]:
            for j in i:
                if len(j) > 0:
                    if j[0] not in tl:
                        tl += [j]
        return tl

    def build_plan(self):
        title = [['I. ', 'II. ', 'III. ', 'IV. ', 'V. ', 'VI. ', 'VII. ', 'VIII. ', 'IX. ', 'X. '],
                    ['A. ', 'B. ', 'C. ', 'D. ', 'E. ', 'F. ', 'G. ', 'H. ', 'I. ', 'J. '],
                    ['1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ', '10. '],
                    ['a. ', 'b. ', 'c. ', 'd. ', 'e. ', 'f. ', 'g. ', 'h. ', 'i. ', 'j. '],
                    ['i. ', 'ii. ', 'iii. ', 'iv. ', 'v. ', 'vi. ', 'vii. ', 'viii. ', 'ix. ', 'x. '],
                    ['->. ', '->. ', '->. ', '->. ', '->. ', '->. ', '->. ', '->. ', '->. ', '->. ']]
        ct = [0, 0, 0, 0, 0, 0]
        built_plan = ['' for k in range(len(self.plan["position"]))]
        
        for p in range(len(self.plan["position"])):
            k = self.plan["position"].index(p)
            for l in self.tag_list:
                if self.plan["ID"][k] == l[1]:
                    order = self.plan["order"][k]
                    built_plan[self.plan["position"][k]] = '___' * order + title[order][ct[order]] + l[0]
                    ct[order] += 1
                    if p < len(self.plan["position"])-1:
                        diff = order - self.plan["order"][self.plan["position"].index(p+1)]
                        if diff > 0:
                            for j in range(diff):
                                ct[order - j] = 0
                    break
        return built_plan

    def move_left_plan(self):
        pos = self.plan_listbox.curselection()
        for j in pos:
            for k in range(len(self.plan["position"])):
                if j == self.plan["position"][k] and self.plan["order"][k] != 0:
                    self.plan["order"][k] -= 1

        self.plan_listbox.delete(0, END)
        new = self.build_plan()
        for k in range(len(new)):
            self.plan_listbox.insert(k, new[k])

        self.save_dict(self.p, 'plan', self.plan)
        self.plan = self.import_dict(self.p, 'plan')
        self.plan_listbox.select_set(pos[0])

    def move_right_plan(self):
        pos = self.plan_listbox.curselection()
        for j in pos:
            for k in range(len(self.plan["position"])):
                if j == self.plan["position"][k] and self.plan["order"][k] != 5:
                    self.plan["order"][k] += 1

        self.plan_listbox.delete(0, END)
        new = self.build_plan()
        for k in range(len(new)):
            self.plan_listbox.insert(k, new[k])

        self.save_dict(self.p, 'plan', self.plan)
        self.plan = self.import_dict(self.p, 'plan')
        self.plan_listbox.select_set(pos[0])

    def move_up_plan(self):
        pos = self.plan_listbox.curselection()[0]
        target = -1
        for k in range(len(self.plan["position"])):
            if pos == self.plan["position"][k]:
                current = k
            elif pos - 1 == self.plan["position"][k]:
                target = k
        if target != -1:
            tp = self.plan["position"][target]
            self.plan["position"][target] = self.plan["position"][current]
            self.plan["position"][current] = tp

        self.plan_listbox.delete(0, END)
        new = self.build_plan()
        for k in range(len(new)):
            self.plan_listbox.insert(k, new[k])

        self.save_dict(self.p, 'plan', self.plan)
        self.plan = self.import_dict(self.p, 'plan')
        if target != -1:
            self.plan_listbox.select_set(pos - 1)
        else:
            self.plan_listbox.select_set(pos) 

    def move_down_plan(self):
        pos = self.plan_listbox.curselection()[0]
        target = -1
        for k in range(len(self.plan["position"])):
            if pos == self.plan["position"][k]:
                current = k
            elif pos + 1 == self.plan["position"][k]:
                target = k
        if target != -1:
            tp = self.plan["position"][target]
            self.plan["position"][target] = self.plan["position"][current]
            self.plan["position"][current] = tp

        self.plan_listbox.delete(0, END)
        new = self.build_plan()
        for k in range(len(new)):
            self.plan_listbox.insert(k, new[k])

        self.save_dict(self.p, 'plan', self.plan)
        self.plan = self.import_dict(self.p, 'plan')
        if target != -1:
            self.plan_listbox.select_set(pos + 1)
        else:
            self.plan_listbox.select_set(pos)

    def add_plan(self, order_option=0):
        # to be called on edit button press
        pos = self.plan_listbox.curselection()
        if len(pos) == 0:
            current_order = 0
            order_option = 0
        else:
            current_order = self.plan["order"][self.plan["position"].index(pos[0])]
        # Get plan tag name
        self.shell_text.delete("1.0", "end-1c")
        self.shell_label.configure(text='New category? :')
        self.window.bind('<Key>', self.next_press)
        self.window.wait_variable(self.var)
        self.window.unbind('<Key>')
        new_tag = self.shell_text.get("1.0", "end-1c")

        # Get tag order
        input_order = current_order + order_option
        if input_order in self.accepted_order:
            if new_tag not in [j[0] for j in self.tag_list]:
                if len(self.plan["ID"]) == 0:
                    new_ID = 1
                else:
                    new_ID = max(self.plan["ID"]) + 1
                self.plan["ID"] += [new_ID]
                # if no position has been selected
                if len(pos) == 0:
                    input_order = 0
                    position = 0
                else:
                    position = pos[0] + 1
                if position in self.plan["position"]:
                    for k in range(len(self.plan["position"])):
                        if position <= self.plan["position"][k]:
                            self.plan["position"][k] += 1
                self.plan["position"] += [position]
                self.plan["order"] += [input_order]
                self.plan["note"] += ['']
                self.blocs["tag"][0] += [[new_tag, new_ID]]

                self.save_dict(self.p, 'blocs', self.blocs)
                self.blocs = self.import_dict(self.p, 'blocs')
                self.tag_list = self.build_tag_list(self.blocs)
                self.save_dict(self.p, 'plan', self.plan)
                self.plan = self.import_dict(self.p, 'plan')

                self.plan_listbox.delete(0, END)
                new = self.build_plan()
                for k in range(len(new)):
                    self.plan_listbox.insert(k, new[k])
                self.plan_listbox.select_set(pos[0]+1)
        self.shell_text.delete("1.0", "end-1c")
        self.shell_label.configure(text='Shell :')
        
    def add_plan_low(self):
        self.add_plan(order_option=1)

    def add_plan_high(self):
        self.add_plan(order_option=-1)

    def delete_plan(self):
        pos = self.plan_listbox.curselection()[0]
        deleted = 0
        for i in range(len(self.plan["position"])):
            k = i - deleted
            if pos < self.plan["position"][k]:
                self.plan["position"][k] -= 1
            elif pos == self.plan["position"][k]:
                deleted = 1
                deleted_id = self.plan["ID"][k]
                del self.plan["position"][k]
                del self.plan["ID"][k]
                del self.plan["order"][k]
                del self.plan["note"][k]

        for k in range(len(self.blocs["tag"])):
            deleted = 0
            for i in range(len(self.blocs["tag"][k])):
                j = i - deleted
                if self.blocs["tag"][k][j][1] == deleted_id:
                    deleted += 1
                    del self.blocs["tag"][k][j]

        self.save_dict(self.p, 'blocs', self.blocs)
        self.blocs = self.import_dict(self.p, 'blocs')
        self.tag_list = self.build_tag_list(self.blocs)
        self.save_dict(self.p, 'plan', self.plan)
        self.plan = self.import_dict(self.p, 'plan')

        self.plan_listbox.delete(0, END)
        new = self.build_plan()
        for k in range(len(new)):
            self.plan_listbox.insert(k, new[k])

    def edit_plan(self):
        pos = self.plan_listbox.curselection()[0]
        # Ask user input
        self.shell_text.delete("1.0", "end-1c")
        old = self.plan_listbox.get(pos)
        self.shell_text.insert(END, old[old.index('. ')+2:])
        self.shell_label.configure(text='Change name to :')

        self.window.bind('<Key>', self.next_press)
        self.window.wait_variable(self.var)
        self.window.unbind('<Key>')
        new_name = self.shell_text.get("1.0", "end-1c")
        self.shell_text.delete("1.0", "end-1c")
        self.shell_label.configure(text='Shell :')

        for k in range(len(self.plan["position"])):
            if pos == self.plan["position"][k]:
                id = self.plan["ID"][k]

        for k in range(len(self.blocs["tag"])):
            for j in range(len(self.blocs["tag"][k])):
                if self.blocs["tag"][k][j][1] == id:
                    self.blocs["tag"][k][j][0] = new_name

        self.tag_list = self.build_tag_list(self.blocs)
        self.plan_listbox.delete(0, END)
        new = self.build_plan()
        for k in range(len(new)):
            self.plan_listbox.insert(k, new[k])

        self.save_dict(self.p, 'blocs', self.blocs)
        self.blocs = self.import_dict(self.p, 'blocs')
        self.tag_list = self.build_tag_list(self.blocs)
        self.save_dict(self.p, 'plan', self.plan)
        self.plan = self.import_dict(self.p, 'plan')

    def edit_notes_from_plan(self):
        self.noting = 1
        self.notes_text.config(bg='#FFFACD')
        self.save_note_but.config(bg = '#CDC9A5')
        # Get corresponding notes
        cursor = self.plan_listbox.curselection()[0]
        pos = self.plan["position"].index(cursor)
        self.plan_listbox.itemconfig(index=cursor, bg = '#CDC9A5')
        self.notes_text.delete('1.0', "end-1c")
        self.notes_text.insert("1.0", self.plan["note"][pos])
        # Wait for save button press
        self.save_note_but.wait_variable(self.save_var)
        note = self.notes_text.get("1.0", "end-1c")
        self.plan["note"][pos] = note

        self.save_dict(self.p, 'plan', self.plan)
        self.plan = self.import_dict(self.p, 'plan')
        self.notes_text.config(bg='white')
        self.plan_listbox.itemconfig(index=cursor, bg = 'white')
        self.save_note_but.config(bg = '#f0f0f0')
        self.noting = 0

    def insert_ref(self):
        cursor = self.notes_text.index(INSERT)
        sources = "("
        selected = self.source_listbox.curselection()
        for k in selected:
            sources += self.source_listbox.get(k) + " ; "
        sources = sources[:-3]
        sources += ") "
        self.notes_text.insert(cursor, sources)

    # Blocs management

    def add_to_blocs(self):
        # Add if not in MultiTagBiblio
        sources, highlights, notes = self.zotero_import()
        for k in range(len(highlights)):
            if highlights[k] == 'Just a note':
                if "NOTE : " + notes[k] not in self.blocs['text']:
                    self.blocs['text'] += ["NOTE : " + notes[k]]
                    self.blocs['source'] += [sources[k]]
                    self.blocs['tag'] += [[]]
            else:
                if highlights[k] not in self.blocs['text']:
                    self.blocs['text'] += [highlights[k]]
                    self.blocs['source'] += [sources[k]]
                    self.blocs['tag'] += [[]]
                if len(notes[k]) > 0:
                    if "COM : " + notes[k] not in self.blocs['text']:
                        self.blocs['text'] += ["COM : " + notes[k]]
                        self.blocs['source'] += [sources[k]]
                        self.blocs['tag'] += [[]]
                        
        
        # Delete if not in zotero
        deleted = 0
        for j in range(len(self.blocs["text"])):
            k = j - deleted
            if self.blocs["text"][k] != None and self.blocs["text"][k] != 'default':
                if "NOTE : " in self.blocs["text"][k]:
                    if self.blocs["text"][k][7:] not in notes:
                        del self.blocs["text"][k]
                        del self.blocs["source"][k]
                        del self.blocs["tag"][k]
                        deleted += 1

                elif "COM : " in self.blocs["text"][k]:
                    if self.blocs["text"][k][6:] not in notes:
                        del self.blocs["text"][k]
                        del self.blocs["source"][k]
                        del self.blocs["tag"][k]
                        deleted += 1
                else:
                    to_delete = 1
                    for p in highlights:
                        if p in self.blocs["text"][k]:
                            to_delete = 0
                            break
                    if to_delete == 1:
                        del self.blocs["text"][k]
                        del self.blocs["source"][k]
                        del self.blocs["tag"][k]
                        deleted += 1

        self.save_dict(self.p, 'blocs', self.blocs)
        self.blocs = self.import_dict(self.p, 'blocs')

        self.source_listbox.delete(0, END)
        for k in unique(self.blocs["source"]):
            self.source_listbox.insert(END, k[0])

    def delete_article(self):
        pos = self.source_listbox.curselection()[0]
        if self.source_listbox.get(pos) != "default":
            article = self.source_listbox.get(pos)
            deleted = 0
            for j in range(len(self.blocs["source"])):
                k = j - deleted
                if self.blocs["source"][k][0] == article:
                    del self.blocs["text"][k]
                    del self.blocs["source"][k]
                    del self.blocs["tag"][k]
                    deleted += 1

            self.save_dict(self.p, 'blocs', self.blocs)
            self.blocs = self.import_dict(self.p, 'blocs')

            self.source_listbox.delete(0, END)
            for k in unique(self.blocs["source"]):
                self.source_listbox.insert(END, k[0])

    def tag_blocs(self):
        self.tagging = 1
        self.tag_but.config(bg = '#CDC9A5')
        self.tag_list = self.build_tag_list(self.blocs)
        selected = []
        source = self.source_listbox.curselection()
        if len(self.blocs_listbox.curselection()) > 0:
            current = self.blocs_listbox.curselection()
            for k in current:
                selected += [self.blocs_listbox.get(k)]

            for k in range(len(current)):
                # Show text in shell
                self.shell_text.delete('1.0', "end-1c")
                self.blocs_listbox.itemconfig(current[k], bg='green')
                self.plan_listbox.selection_clear(0, END)
                add_tag = self.blocs["text"].index(selected[k])
                self.shell_text.insert(END, self.blocs["text"][add_tag])
                # Show existing tags
                # Reset colors
                for p in self.plan["position"]:
                    self.plan_listbox.itemconfig(p, bg='white')
                # Color existing tags
                for j in range(len(self.blocs["text"])):
                    if self.blocs["text"][j] == self.blocs["text"][add_tag]:
                        for i in self.blocs["tag"][j]:
                            for m in range(len(self.plan["ID"])):
                                if i[1] == self.plan["ID"][m]:
                                    self.plan_listbox.select_set(self.plan["position"][m])
                                    # switch to select set
                # Wait for button press
                self.window.bind('<Key>', self.next_press)
                self.window.wait_variable(self.var)
                self.window.unbind('<Key>')
                self.blocs_listbox.itemconfig(current[k], bg='white')
                add = []
                focus = self.plan_listbox.curselection()
                # if a tag selection has been made
                if len(focus) > 0:
                    for m in focus:
                        tag = [m]
                        # To associate bloc to selected tag but also his parents.
                        for o in range(self.plan["order"][self.plan["position"].index(tag[0])]):
                            tag += [self.get_parent(tag[-1])]
                        for i in tag:
                            id = self.plan["ID"][self.plan["position"].index(i)]
                            for j in self.tag_list:
                                if j[1] == id:
                                    add += [j]

                    self.blocs["tag"][add_tag] = add

        else:
            sep = self.source_listbox.get(source)
            for k in range(len(self.blocs["text"])):
                if self.blocs["source"][k][0] == sep:
                    selected += [self.blocs["text"][k]]

            deleted = 0
            for l in range(len(selected)):
                k = l - deleted
                # Color current focus block
                self.blocs_listbox.itemconfig(l, bg='green')
                # Show text in shell
                self.shell_text.delete('1.0', "end-1c")

                if self.merge_var.get() == 0:
                    add_tag = self.blocs["text"].index(selected[k])

                if k < len(selected) - 1:
                    next = self.blocs["text"].index(selected[k + 1])
                    self.shell_text.insert(END, self.blocs["text"][add_tag] + '\n\nNext : \n\n' + self.blocs["text"][next])
                else:
                    self.shell_text.insert(END, self.blocs["text"][add_tag])
                
                self.merge_var.set(0)

                # Show existing tags
                # Reset colors
                for r in self.plan["position"]:
                    self.plan_listbox.itemconfig(r, bg='white')
                # Color existing tags
                for j in range(len(self.blocs["text"])):
                    if self.blocs["text"][j] == self.blocs["text"][add_tag]:
                        for i in self.blocs["tag"][j]:
                            for m in range(len(self.plan["ID"])):
                                if i[1] == self.plan["ID"][m]:
                                    self.plan_listbox.itemconfig(self.plan["position"][m], bg='green')
                # Wait for button press
                self.window.bind('<Key>', self.next_press)
                self.window.wait_variable(self.var)
                self.window.unbind('<Key>')
                if self.merge_var.get() == 1:
                    deleted += 1
                    # Merge text
                    self.blocs["text"][add_tag] += self.blocs["text"][next]
                    # Merge existing tags
                    self.blocs["tag"][add_tag] += self.blocs["tag"][next]
                    del selected[k + 1]
                    del self.blocs["text"][next]
                    del self.blocs["source"][next]
                    del self.blocs["tag"][next]

                    self.save_dict(self.p, 'blocs', self.blocs)
                    self.blocs = self.import_dict(self.p, 'blocs')

                add = []
                focus = self.plan_listbox.curselection()
                # if a tag selection has been made
                if len(focus) > 0 :
                    for m in focus:
                        tag = [m]
                        # To associate bloc to selected tag but also his parents.
                        for k in range(self.plan["order"][self.plan["position"].index(tag[0])]):
                            tag += [self.get_parent(tag[-1])]
                        for i in tag:
                            id = self.plan["ID"][self.plan["position"].index(i)]
                            for j in self.tag_list:
                                if j[1] == id:
                                    add += [j]

                    self.blocs["tag"][add_tag] += add
                    unique_tags = []
                    for n in self.blocs["tag"][add_tag]:
                        if n not in unique_tags:
                            unique_tags += [n]
                    self.blocs["tag"][add_tag] = unique_tags

        self.save_dict(self.p, 'blocs', self.blocs)
        self.blocs = self.import_dict(self.p, 'blocs')
        self.tag_list = self.build_tag_list(self.blocs)
        self.save_dict(self.p, 'plan', self.plan)
        self.plan = self.import_dict(self.p, 'plan')

        self.tagging = 0
        self.tag_but.config(bg = '#f0f0f0')

    def get_parent(self, position):
        parent = []
        for k in range(len(self.plan["position"])):
            if self.plan["position"][k] < position and self.plan["order"][k] == self.plan["order"][self.plan["position"].index(position)] - 1:
                parent += [self.plan["position"][k]]
        return max(parent)

    def next_press(self, event):
        if event.keysym == 'Return':
            self.var.set(1)

    # Blocs visualization filters

    def access_from_plan(self):
        # to be called as soon as pressed
        # get selected from list
        position_selected = 0
        blocs_selected = []
        for k in range(len(self.blocs["tag"])):
            for l in range(len(self.blocs["tag"][k])):
                if self.plan["ID"][position_selected] == self.blocs["tag"][k][l]:
                    blocs_selected += [self.blocs["text"][k]]

    def blocs_filter_plan(self, event):
        if self.tagging == 0:
            if self.noting == 0:
                self.notes_text.delete("1.0", "end-1c")
            cursor = self.plan_listbox.curselection()
            for i in cursor:
                pos = self.plan["position"].index(i)
                selected = []
                sources = []
                # Skip the firt 'default' source
                for k in range(1, len(self.blocs["text"])):
                    for l in range(len(self.blocs["tag"][k])):
                        if len(self.blocs["tag"][k][l]) > 0:
                            if self.blocs["tag"][k][l][1] == self.plan["ID"][pos]:
                                selected += [self.blocs["text"][k]]
                                sources += [self.blocs["source"][k]]
                sources = unique(sources)

                self.blocs_listbox.delete(0, END)
                for k in [unique(self.blocs["source"]).index(l) for l in self.blocs["source"]]:
                    self.source_listbox.itemconfig(k, bg="white")
                for k in selected:
                    self.blocs_listbox.insert(END, k)
                for k in sources:
                    index = self.source_listbox.get(0, "end").index(k[0])
                    self.source_listbox.itemconfig(index, bg='green')
                
                if self.noting == 0:
                    self.notes_text.insert(END, self.plan["note"][pos] + '\n----\n')

    def blocs_filter_sources(self, event):
        pos = self.source_listbox.curselection()[0]
        selected = []
        for k in range(len(self.blocs["text"])):
            if self.blocs["source"][k][0] == self.source_listbox.get(pos):
                selected += [self.blocs["text"][k]]

        self.blocs_listbox.delete(0, END)
        for k in selected:
            self.blocs_listbox.insert(END, k)

    def blocs_filter_search(self):
        self.blocs_listbox.delete(0, END)
        request = self.search_text.get("1.0", "end-1c")
        for k in self.blocs["text"]:
            if request.lower() in k.lower():
                self.blocs_listbox.insert(END, k)

    def blocs_main_subjects(self):
        tensors = []
        # load pre-trained model
        model = SentenceTransformer('all-MiniLM-L6-v2')
        for k in self.blocs["text"]:
            if k != 'default':
                selected = model.encode(k)
                compare = model.encode(self.blocs["text"])
                cos_sim = util.cos_sim(selected, compare)
                tensors += [cos_sim.tolist()[0]]

        # Classifying vectorial products for each bloc
        linkage_data = linkage(tensors, method='ward', metric='euclidean')
        dendrogram(linkage_data)

        self.shell_text.delete('1.0', "end-1c")
        self.shell_label.configure(text='Number of clusters? :')
        plt.show()
        self.window.bind('<Key>', self.next_press)
        self.window.wait_variable(self.var)
        self.window.unbind('<Key>')
        nb_clusters = int(self.shell_text.get("1.0", "end-1c"))

        hierarchical_cluster = AgglomerativeClustering(n_clusters=nb_clusters, affinity='euclidean', linkage='ward')
        labels = hierarchical_cluster.fit_predict(tensors)
        groups = ['' for k in unique(labels)]

        for k in range(len(self.blocs["text"])-1):
            groups[labels[k]] += self.blocs["text"][k+1]

        self.shell_text.delete('1.0', "end-1c")
        for k in range(len(groups)):
            # Frequent keywords and word combinations in blocs
            words = nltk.tokenize.word_tokenize(groups[k])
            stop_words = set(nltk.corpus.stopwords.words("english"))
            stop_words.update(',', '.', '(', ')', '[', ']', 'Figure')
            filtered_words = []
            # Filter words and set to their lemma
            lemmatizer = nltk.stem.WordNetLemmatizer()
            for word in words:
                if word.casefold() not in stop_words:
                    filtered_words.append(lemmatizer.lemmatize(word))
            self.shell_text.insert(END, "Group " + str(k) + " :\n")
            for k in nltk.FreqDist(filtered_words).most_common(2):
                self.shell_text.insert(END, str(k)+"\n")

            # Study bi-collocation information
            bigram_measures = nltk.collocations.BigramAssocMeasures()
            finder = nltk.collocations.BigramCollocationFinder.from_words(filtered_words)
            for k in finder.nbest(bigram_measures.likelihood_ratio, 2):
                self.shell_text.insert(END, str(k) + "\n")

    def read_blocs(self, event):
        if self.tagging == 0:
            # Reset widgets
            self.shell_text.delete('1.0', "end-1c")
            for k in [unique(self.blocs["source"]).index(l) for l in self.blocs["source"]]:
                self.source_listbox.itemconfig(k, bg="white")
            for k in self.plan["position"]:
                if self.plan_listbox.itemcget(k, "bg") != '#CDC9A5':
                    self.plan_listbox.itemconfig(k, bg='white')

            for i in self.blocs_listbox.curselection():
                # Display in shell
                self.shell_text.insert(END, self.blocs_listbox.get(i))

                # Show corresponding tags and sources
                for j in range(len(self.blocs["text"])):
                    if self.blocs["text"][j] == self.blocs_listbox.get(i):
                        source = self.blocs["source"][j]
                        self.source_listbox.itemconfig(unique(self.blocs["source"]).index(source), bg='green')

                        for k in self.blocs["tag"][j]:
                            for m in range(len(self.plan["ID"])):
                                if k[1] == self.plan["ID"][m]:
                                    if self.plan_listbox.itemcget(self.plan["position"][m], "bg") != '#CDC9A5':
                                        self.plan_listbox.itemconfig(self.plan["position"][m], bg='green')

                self.shell_text.insert(END, '\n----\n')

    def send_key(self):
        pos = self.source_listbox.curselection()[0]
        for k in self.blocs["source"]:
            if k[0] == self.source_listbox.get(pos):
                key = k[1]
                break
        ArticleInfo(key=key, zotero=self.zotero)

    # Exchanges with Zotero

    def locate_zotero(self):
        path = askdirectory(parent=self.window, title='Select Zotero folder location')
        target_collection = tk.simpledialog.askstring(
            "Zotero collection", "Enter the parent Zotero collection you are working with : ")
        d = dict(path=path, target_collection=target_collection)
        with open(p + "\\Zotero_data.pkl", 'wb') as f:
            pickle.dump(d, f)
        
        self.zotero = self.import_dict(self.p, 'Zotero_data')
        self.zotero_path.set(self.zotero['path'] if self.zotero else 'Not yet connected to Zotero.')
        return

    def zotero_import(self):
        # standard database connection
        conn = sqlite3.connect(self.zotero['path'] + '/zotero.sqlite')
        cur = conn.cursor()
        
        # IF YOU WANT TO EXPLORE ZOTERO DATABASE TABLES NAMES
        # cursor = cur.execute('SELECT * FROM collections')
        # names = list(map(lambda x: x[0], cursor.description))
        # print(names)


        sql = """SELECT collectionID FROM collections WHERE collectionName='""" + self.zotero["target_collection"] + "';"
        cursor = cur.execute(sql)
        parent = cursor.fetchall()[0][0]

        sql = """SELECT collectionID FROM collections WHERE parentCollectionID=""" + str(parent) + ";"
        cursor = cur.execute(sql)
        childs = [k[0] for k in cursor.fetchall()]

        familly = [parent]
        while len(childs) != 0:
            tp = []
            for k in childs:
                familly += [k]
                sql = """SELECT collectionID FROM collections WHERE parentCollectionID=""" + str(k) + ";"
                cursor = cur.execute(sql)
                tp += [k[0] for k in cursor.fetchall()]
            childs = tp

        items = []
        for k in familly:
            sql = """SELECT itemID FROM collectionItems WHERE collectionID=""" + str(k) + ";"
            cursor = cur.execute(sql)
            items += [k[0] for k in cursor.fetchall()]
        items = unique(items)

        sql = """SELECT parentItemID, type, text, comment, sortIndex FROM itemAnnotations WHERE type=1 OR type=2"""
        cursor = cur.execute(sql)
        all_annotations = cursor.fetchall()

        items_annotations = [[] for k in items]
        for k in all_annotations:
            sql = """SELECT parentItemID FROM itemAttachments WHERE itemID=""" + str(k[0]) + ";"
            cursor = cur.execute(sql)
            result = cursor.fetchall()[0][0]
            if result in items:
                pos = items.index(result)
                items_annotations[pos] += [k]
        
        source = [[] for k in items]
        for k in range(len(items)):
            sql = """SELECT creatorID FROM itemCreators WHERE itemID=""" + str(items[k]) + " AND orderIndex=0;"
            cursor = cur.execute(sql)
            creator_ID = cursor.fetchall()

            sql = """SELECT valueID FROM itemData WHERE itemID=""" + str(items[k]) + " AND fieldID=6;"
            cursor = cur.execute(sql)
            date_ID = cursor.fetchall()

            if len(creator_ID) > 0:
                sql = """SELECT lastName FROM creators WHERE creatorID=""" + str(creator_ID[0][0]) + ";"
                cursor = cur.execute(sql)
                first_author = cursor.fetchall()[0][0]
            else:
                first_author = '?'

            if len(date_ID) > 0 :
                sql = """SELECT value FROM itemDataValues WHERE valueID=""" + str(date_ID[0][0]) + ";"
                cursor = cur.execute(sql)
                date = cursor.fetchall()[0][0][:4]

            else:
                date = '?'
            
            source[k] = [first_author + ' et al., ' + date, items[k]]

        sources, highlights, notes = [], [], []
        for k in range(len(items_annotations)):
            items_annotations[k] = sorted(list(items_annotations[k]), key=lambda d: d[4])

            for j in range(len(items_annotations[k])):
                if items_annotations[k][j][2] != None and items_annotations[k][j][3] == None:
                    highlights += [items_annotations[k][j][2]]
                    notes += ['']
                    sources += [source[k]]
                elif items_annotations[k][j][2] == None and items_annotations[k][j][3] != None:
                    highlights += ['Just a note']
                    notes += [items_annotations[k][j][3]]
                    sources += [source[k]]
                elif items_annotations[k][j][2] != None and items_annotations[k][j][3] != None:
                    highlights += [items_annotations[k][j][2]]
                    notes += [items_annotations[k][j][3]]
                    sources += [source[k]]

        return sources, highlights, notes

    # Exchanges with word

    def export_all(self):
        cwd = os.getcwd()
        doc = docx.Document()

        headers = [[] for k in self.tag_list]
        for k in range(len(self.tag_list)):
            for m in range(len(self.plan["ID"])):
                if self.tag_list[k][1] == self.plan["ID"][m]:
                    sources = []
                    # Skip the firt 'default' source
                    for n in range(1, len(self.blocs["text"])):
                        for o in range(len(self.blocs["tag"][n])):
                            if len(self.blocs["tag"][n][o]) > 0:
                                if self.blocs["tag"][n][o][1] == self.plan["ID"][m]:
                                    sources += [self.blocs["source"][n][0]]
                    printed_sources = ''
                    for p in unique(sources):
                        printed_sources += p + " ; "
                    headers[self.plan["position"][m]] = [self.tag_list[k][0], self.plan["order"][m], self.plan["note"][m], printed_sources[:-3]]

        
        for k in headers:
            if k != []:
                doc.add_heading(k[0], level=k[1]+1)
                if len(k[2]) > 0:
                    doc.add_paragraph(k[2])
                    doc.add_paragraph(str(k[3]))

        filepath = cwd + "\\docx\\biblio_analysis.docx"
        doc.save(filepath)

        # for windows users only
        os.startfile(filepath)

    # Backup files and version in case of bug

    def backup(self):
        cwd = os.getcwd()
        folder = datetime.today().strftime('%Y-%m-%d-%H-%M')
        os.makedirs(cwd + "\\Backup\\" + folder)

        shutil.copyfile(cwd + "\\MultiTagBiblio.py", cwd + "\\Backup\\" + folder + "\\MultiTagBiblio.py")
        shutil.copytree(cwd + "\\Storage", cwd + "\\Backup\\" + folder + "\\Storage")



class ArticleInfo:
    def __init__(self, key, zotero):
        win2 = Tk()
        win2.title('Article metadata')
        win2.geometry("600x400")
        my_label = HTMLLabel(win2, height=20, width=70, html=self.get_meta(key, zotero))
        my_label.place(x = 10, y = 10)
        win2.mainloop()

    def get_meta(self, key, zotero):
        # standard database connection
        conn = sqlite3.connect(zotero['path'] + '/zotero.sqlite')
        cur = conn.cursor()

        sql = """SELECT valueID FROM itemData WHERE itemID=""" + str(key) + " AND fieldID=1;"
        cursor = cur.execute(sql)
        title_ID = cursor.fetchall()

        if len(title_ID) > 0 :
            sql = """SELECT value FROM itemDataValues WHERE valueID=""" + str(title_ID[0][0]) + ";"
            cursor = cur.execute(sql)
            title = cursor.fetchall()[0][0]
        else:
            title = ""

        sql = """SELECT valueID FROM itemData WHERE itemID=""" + str(key) + " AND fieldID=6;"
        cursor = cur.execute(sql)
        date_ID = cursor.fetchall()

        if len(date_ID) > 0 :
            sql = """SELECT value FROM itemDataValues WHERE valueID=""" + str(date_ID[0][0]) + ";"
            cursor = cur.execute(sql)
            date = cursor.fetchall()[0][0]
        else:
            date = ""

        sql = """SELECT creatorID, orderIndex FROM itemCreators WHERE itemID=""" + str(key) + ";"
        cursor = cur.execute(sql)
        creator_ID = sorted(cursor.fetchall(), key=lambda d: d[1])

        authors = ""
        for k in creator_ID:
            sql = """SELECT lastName FROM creators WHERE creatorID=""" + str(k[0]) + ";"
            cursor = cur.execute(sql)
            authors += cursor.fetchall()[0][0] + " "
            sql = """SELECT firstName FROM creators WHERE creatorID=""" + str(k[0]) + ";"
            cursor = cur.execute(sql)
            authors += cursor.fetchall()[0][0] + ", "


        sql = """SELECT valueID FROM itemData WHERE itemID=""" + str(key) + " AND fieldID=2;"
        cursor = cur.execute(sql)
        abstract_ID = cursor.fetchall()

        if len(abstract_ID) > 0 :
            sql = """SELECT value FROM itemDataValues WHERE valueID=""" + str(abstract_ID[0][0]) + ";"
            cursor = cur.execute(sql)
            abstract = cursor.fetchall()[0][0]
        else:
            abstract = ""

        return (title + '<br>-<br>'
                + authors + '<br>-<br>'
                + date[:4] + '<br>-<br>'
                + abstract)


def init_dict():
    cwd = os.getcwd()
    p = cwd + "\\Storage"
    # Check whether the specified path exists or not
    if not os.path.exists(p):
        os.makedirs(p)
        d = dict(text=['default'], source=[['default', '']], tag=[[]])
        with open(p + "\\blocs.pkl", 'wb') as f:
            pickle.dump(d, f)

        d = dict(position=[], ID=[], order=[], note=[])
        with open(p + "\\plan.pkl", 'wb') as f:
            pickle.dump(d, f)

        with open(p + "\\Zotero_data.pkl", 'wb') as f:
            pickle.dump({}, f)
        """
        print('No Zotero folder found')
        dialog = Tk()
        dialog.withdraw()
        path = askdirectory(parent=dialog, title='Select Zotero folder location')
        dialog.destroy()
        target_collection = input("Enter the parent Zotero collection you are working with : ")
        d = dict(path=path, target_collection=target_collection)
        with open(p + "\\Zotero_data.pkl", 'wb') as f:
            pickle.dump(d, f)
        """
    if not os.path.exists(cwd + "\\docx"):
        os.makedirs(cwd + "\\docx")
    if not os.path.exists(cwd + "\\Backup"):
        os.makedirs(cwd + "\\Backup")
    
    # Create windows system shortcuts
    startup = "C:\\Users\\" + os.getlogin() + "\\AppData\\Roaming\\Microsoft\\Windows\\Start Menu\\Programs"
    target = cwd + "\\MultiTagBiblio.bat"
    if not os.path.exists(startup + "\\MultiTagBiblio.lnk"):
        ask_user = input("Do you want to create a shortcut ? (y/n) : ")
        if ask_user == "y":
            print("Creating quick access shortcuts")
            shell = win32com.client.Dispatch("WScript.Shell")
            shortcut = shell.CreateShortCut(startup + "\\MultiTagBiblio.lnk")
            shortcut.Targetpath = 'cmd.exe'
            shortcut.Arguments = '/C "' + target + '"'
            shortcut.WorkingDirectory = cwd
            shortcut.IconLocation = cwd + "\\bin\\MTB_logo.ico"
            shortcut.WindowStyle = 1
            shortcut.save()


def unique(X):
    only_one = []
    for k in X:
        if k not in only_one:
            only_one += [k]
    return only_one


# Script

myappid = 'inrae.multitagbiblio.zotero.1.0' # arbitrary string
ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)

BG_COLOR = '#96CDCD'

init_dict()
win = Tk()
win.title('Multi Tag Biblio')
win.state('zoomed')
win.configure(bg=BG_COLOR, bd=10)
win.iconbitmap("bin\\MTB_logo.ico")
mt = Biblio(win)
win.mainloop()
